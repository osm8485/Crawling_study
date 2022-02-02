from docx import Document
from bs4 import BeautifulSoup
import requests
import time
import pyautogui

# 사용자입력
keyword = pyautogui.prompt('검색어를 입력하세요 : ')
lastpage = int(pyautogui.prompt('몇페이지까지 크롤링 할까요? : '))

# 워드문서 생성하기
document = Document()

pageNum = 1
for page in range(1, lastpage*10, 10):
    print(f"=================={pageNum}페이지 입니다==================")
    response = requests.get(
        f'https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}&start={page}')
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')
    articles = soup.select('div.info_group')  # 뉴스 기사 div 10개 추출

    for article in articles:
        links = article.select('a.info')  # 리스트
        if len(links) >= 2:  # 링크가 2개 이상이면
            url = links[1].attrs['href']  # 두번째 링크의 href를 추출
            response = requests.get(url, headers={'User-Agent': 'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            # 만약 연예뉴스라면
            if "entertain" in response.url:  # 그냥 url사용시 리다이렉션으로 NoneType오류 또 발생, response.url로사용!
                title = soup.select_one('.end_tit')
                content = soup.select_one('#articeBody')
            elif "sports" in response.url:
                title = soup.select_one('h4.title')
                # content의 태그가 필요하므로 .text() 사용x
                content = soup.select_one('#newsEndContents')
                # 불필요한 div, p태그 삭제
                divs = content.select('div')
                paragraphs = content.select('p')
                for div in divs:
                    div.decompose()
                for p in paragraphs:
                    p.decompose()
            else:
                title = soup.select_one('#articleTitle')
                content = soup.select_one('#articleBodyContents')
            print("==============링크=============\n", url.strip())
            print("==============제목=============\n", title.text.strip())
            print("==============본문=============\n", content.text.strip())
            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())
            document.save(f"C:/Users/o9707/Desktop/대학교/inflearn/crawling/실전편/01.네이버뉴스/{keyword}_test.docx")
            time.sleep(0.3)
    pageNum += 1

# 워드 문서 저장하기
#document.save(f"C:/Users/o9707/Desktop/대학교/inflearn/crawling/실전편/01.네이버뉴스/{keyword}_test.docx")